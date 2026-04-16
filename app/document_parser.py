"""
document_parser.py — MVP-2b

Thin wrappers around pypdf + python-docx + stdlib email that turn uploaded
document bytes into plain text. The text then feeds into
`copilot_engine.extract_demand`, which runs the Claude-Haiku item extractor
and catalog matcher.

Scope is deliberately narrow:
- PDF: visible text layer only (no OCR — scanned PDFs return empty)
- DOCX: paragraph + table text (merged with newlines)
- EML / raw email: body text via stdlib `email.parser`
- TXT / CSV: decoded as UTF-8 with best-effort fallback

Keep it pure-Python (no system deps) so Railway's Dockerfile stays fast.
"""
from __future__ import annotations

import io
import logging
from email import policy
from email.parser import BytesParser

log = logging.getLogger(__name__)


def detect_format(filename: str, content_type: str, raw: bytes) -> str:
    """Best-effort format detection: filename ext wins, content-type is
    tie-breaker, magic bytes for the common cases."""
    name = (filename or "").lower()
    ct = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ct or raw[:4] == b"%PDF":
        return "pdf"
    if name.endswith(".docx") or "wordprocessingml" in ct or "officedocument" in ct:
        return "docx"
    if name.endswith(".eml") or ct.startswith("message/"):
        return "eml"
    if name.endswith(".txt") or name.endswith(".csv") or ct.startswith("text/"):
        return "text"
    # Zip starts with PK — could be docx if ext was lost
    if raw[:2] == b"PK":
        return "docx"
    return "unknown"


# Threshold below which we consider the PDF text layer empty / useless and
# fall back to OCR. Scanned invoices typically return 0–20 chars.
_PDF_TEXT_MIN_CHARS = 50


def extract_text_from_pdf(raw: bytes) -> str:
    """Pull text from a PDF. Tries the native text layer first (cheap via
    pypdf). If the file is a scan (no text layer), falls back to Tesseract
    OCR through pdf2image — configurable via settings.pdf_ocr_enabled."""
    try:
        from pypdf import PdfReader
    except ImportError:
        log.error("pypdf not installed — cannot parse PDF")
        return ""

    text = ""
    try:
        reader = PdfReader(io.BytesIO(raw))
        parts = []
        for page in reader.pages[:30]:  # cap at 30 pages
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt)
        text = "\n\n".join(parts).strip()
    except Exception as exc:
        log.warning("PDF text-layer parse failed: %s", exc)

    if len(text) >= _PDF_TEXT_MIN_CHARS:
        return text

    # Fallback to OCR
    ocr_text = _ocr_pdf(raw)
    if ocr_text and len(ocr_text) > len(text):
        log.info("PDF OCR recovered %d chars (text layer had %d)", len(ocr_text), len(text))
        return ocr_text
    return text


def _ocr_pdf(raw: bytes) -> str:
    """Rasterise each page via pdf2image and run Tesseract (pol+eng).

    Bails out (returns empty string) when any of the optional deps are
    missing or the Tesseract binary is not on PATH — callers should treat
    this as "no OCR available" and keep the empty text-layer string.
    """
    try:
        from app.config import settings
        if not getattr(settings, "pdf_ocr_enabled", True):
            return ""
    except Exception:
        pass  # config import fail shouldn't block OCR

    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError as exc:
        log.warning("OCR deps not installed: %s", exc)
        return ""

    try:
        images = convert_from_bytes(raw, dpi=200, first_page=1, last_page=10)
    except Exception as exc:
        log.warning("pdf2image rasterisation failed: %s", exc)
        return ""

    parts: list[str] = []
    for img in images:
        try:
            text = pytesseract.image_to_string(img, lang="pol+eng")
        except pytesseract.TesseractNotFoundError:
            log.warning("Tesseract binary not installed on PATH — skipping OCR")
            return ""
        except Exception as exc:
            log.warning("Tesseract OCR on page failed: %s", exc)
            continue
        if text and text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def extract_text_from_docx(raw: bytes) -> str:
    """Pull paragraphs + table cells from a DOCX, merged with newlines."""
    try:
        from docx import Document
    except ImportError:
        log.error("python-docx not installed — cannot parse DOCX")
        return ""

    try:
        doc = Document(io.BytesIO(raw))
    except Exception as exc:
        log.warning("DOCX parse failed: %s", exc)
        return ""

    chunks: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def extract_text_from_eml(raw: bytes) -> str:
    """Pull From/Subject + plain-text body from an .eml file.

    Prefers `text/plain` parts; falls back to stripped `text/html` so the
    LLM still gets something to work on.
    """
    try:
        msg = BytesParser(policy=policy.default).parsebytes(raw)
    except Exception as exc:
        log.warning("EML parse failed: %s", exc)
        return _decode_bytes(raw)

    header_lines = []
    for key in ("From", "Subject"):
        val = msg.get(key)
        if val:
            header_lines.append(f"{key}: {val}")
    body = ""
    if msg.is_multipart():
        # Walk and grab first text/plain; fall back to html if none found
        plain = None
        html = None
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain" and plain is None:
                plain = _safe_get_content(part)
            elif ctype == "text/html" and html is None:
                html = _safe_get_content(part)
        body = plain or html or ""
    else:
        body = _safe_get_content(msg)

    return "\n".join(filter(None, ["\n".join(header_lines), body])).strip()


def _safe_get_content(part) -> str:
    try:
        return part.get_content() or ""
    except Exception:
        try:
            payload = part.get_payload(decode=True) or b""
            return _decode_bytes(payload)
        except Exception:
            return ""


def _decode_bytes(raw: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1250", "iso-8859-2", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_text(filename: str, content_type: str, raw: bytes) -> tuple[str, str]:
    """Dispatcher — returns (extracted_text, detected_format).

    Format is returned so callers can surface the choice to the user
    ('Wykryto PDF, 12 stron...') and for telemetry.
    """
    fmt = detect_format(filename, content_type, raw)
    if fmt == "pdf":
        return extract_text_from_pdf(raw), "pdf"
    if fmt == "docx":
        return extract_text_from_docx(raw), "docx"
    if fmt == "eml":
        return extract_text_from_eml(raw), "eml"
    if fmt == "text":
        return _decode_bytes(raw), "text"
    # Last resort — try decoding as text so LLM still has a shot
    return _decode_bytes(raw), "unknown"
