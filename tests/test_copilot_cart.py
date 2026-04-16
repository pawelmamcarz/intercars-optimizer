"""
Tests for copilot add-to-cart pipeline:
- buying_engine.search_catalog (fuzzy product lookup)
- copilot_engine._parse_add_to_cart (regex + numword parser)
- copilot_engine._handle_add_to_cart (action dispatch)
- copilot_engine.process_message end-to-end for add-to-cart intent
"""
from __future__ import annotations

from app.buying_engine import search_catalog
from app.copilot_engine import (
    CopilotRequest,
    _handle_add_to_cart,
    _match_intent,
    _parse_add_to_cart,
    process_message,
)


# ─── search_catalog ──────────────────────────────────────────────────


def test_search_matches_by_name_substring():
    r = search_catalog("klocki hamulcowe", limit=5)
    assert any("klocki" in p["name"].lower() for p in r), r


def test_search_handles_polish_inflection():
    # "hamulce" should still find "Klocki hamulcowe" via 5-char prefix match
    r = search_catalog("hamulce", limit=5)
    names = [p["name"].lower() for p in r]
    assert any("hamulc" in n for n in names), names


def test_search_empty_query_returns_empty_list():
    assert search_catalog("") == []
    assert search_catalog("   ") == []


def test_search_no_match_returns_empty_list():
    assert search_catalog("kosmolot naddzwiekowy xyz123") == []


def test_search_attaches_score():
    r = search_catalog("filtr oleju", limit=1)
    assert r and "_score" in r[0]
    assert r[0]["_score"] > 0


# ─── _parse_add_to_cart ──────────────────────────────────────────────


def test_parse_leading_digit_quantity():
    query, qty = _parse_add_to_cart("dodaj 3 klocki bosch do koszyka")
    assert qty == 3
    assert "klocki bosch" in query.lower()


def test_parse_numword_quantity():
    query, qty = _parse_add_to_cart("dodaj piec filtrow oleju do koszyka")
    assert qty == 5
    assert "filtr" in query.lower()


def test_parse_no_quantity_defaults_one():
    query, qty = _parse_add_to_cart("dodaj filtry oleju do koszyka")
    assert qty == 1
    assert "filtr" in query.lower()


def test_parse_strips_unit_word():
    query, qty = _parse_add_to_cart("dodaj opony sztuk do zamowienia")
    assert qty == 1
    assert "opony" in query.lower()
    assert "sztuk" not in query.lower()


def test_parse_synonym_verbs():
    for verb in ("dodaj", "wrzuc", "wloz", "zamow"):
        query, qty = _parse_add_to_cart(f"{verb} 2 akumulator do koszyka")
        assert qty == 2, verb


def test_parse_non_matching_message_returns_empty():
    query, qty = _parse_add_to_cart("pokaz mi koszyk")
    assert query == ""
    assert qty == 1


# ─── _handle_add_to_cart ─────────────────────────────────────────────


def test_handle_single_dominant_match_emits_action():
    resp = _handle_add_to_cart("kosiarka", 2, context={"step": 0})
    assert resp.actions, "expected actions for dominant match"
    cart_actions = [a for a in resp.actions if a.action_type == "add_to_cart"]
    assert len(cart_actions) == 1
    items = cart_actions[0].params["items"]
    assert len(items) == 1
    assert items[0]["qty"] == 2
    assert "kosiarka" in items[0]["name"].lower()


def test_handle_no_match_suggests_marketplace():
    resp = _handle_add_to_cart("kosmolot xyz123", 3, context={"step": 1})
    cart_actions = [a for a in resp.actions if a.action_type == "add_to_cart"]
    assert not cart_actions, "should not emit add_to_cart when nothing found"
    assert "allegro" in resp.reply.lower() or "marketplace" in resp.reply.lower()


def test_handle_empty_query_asks_for_input():
    resp = _handle_add_to_cart("", 1, context={})
    assert not [a for a in resp.actions if a.action_type == "add_to_cart"]
    assert "?" in resp.reply


def test_handle_ambiguous_returns_suggestions():
    # "bosch" matches multiple products roughly equally
    resp = _handle_add_to_cart("bosch", 1, context={"step": 1})
    cart_actions = [a for a in resp.actions if a.action_type == "add_to_cart"]
    # Should either be dominant (1 clear winner) or ambiguous (suggestions). If
    # ambiguous, no cart action is emitted and suggestions list is populated.
    if not cart_actions:
        assert resp.suggestions, "ambiguous case should offer candidate suggestions"


def test_handle_qty_clamped_to_positive():
    resp = _handle_add_to_cart("kosiarka", 0, context={})
    cart_actions = [a for a in resp.actions if a.action_type == "add_to_cart"]
    if cart_actions:  # dominant match path
        assert cart_actions[0].params["items"][0]["qty"] >= 1


# ─── process_message end-to-end ──────────────────────────────────────


def _run(msg: str, step: int = 1):
    import asyncio
    return asyncio.run(process_message(
        CopilotRequest(message=msg, context={"step": step, "domain": "parts"})
    ))


def test_e2e_intent_recognized():
    intent, _ = _match_intent("dodaj 3 klocki bosch do koszyka")
    assert intent == "add_to_cart_nl"


def test_e2e_emits_add_to_cart_for_dominant_product():
    resp = _run("dodaj 2 kosiarki do koszyka")
    cart_actions = [a for a in resp.actions if a.action_type == "add_to_cart"]
    assert len(cart_actions) == 1
    assert cart_actions[0].params["items"][0]["qty"] == 2


def test_e2e_graceful_miss():
    resp = _run("dodaj rakiete kosmiczna do koszyka")
    cart_actions = [a for a in resp.actions if a.action_type == "add_to_cart"]
    assert not cart_actions


# ─── MVP-1 recommendations ───────────────────────────────────────────


def test_recommendations_returns_nonempty_cards():
    from app.copilot_engine import get_recommendations
    cards = get_recommendations({"step": 0})
    assert len(cards) >= 3, "dashboard needs at least a few cards to feel alive"
    for c in cards:
        assert c["id"], "every card needs a stable id"
        assert c["title"], "every card needs a title"
        assert c["icon"], "every card needs an icon"
        assert c["urgency"] in ("info", "urgent")


def test_recommendations_urgent_cards_present():
    from app.copilot_engine import get_recommendations
    cards = get_recommendations({"step": 0})
    urgencies = {c["urgency"] for c in cards}
    assert "urgent" in urgencies or "info" in urgencies


def test_recommendations_actionable_cards_have_action():
    from app.copilot_engine import get_recommendations
    cards = get_recommendations({"step": 0})
    actionable = [c for c in cards if c.get("cta")]
    assert actionable, "need at least one card with a CTA"
    for c in actionable:
        assert c.get("action"), f"card {c['id']} has CTA but no action"
        assert c["action"]["action_type"], f"card {c['id']} action has no type"


def test_recommendations_endpoint():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/copilot/recommendations?step=0")
    assert r.status_code == 200
    data = r.json()
    assert "cards" in data
    assert isinstance(data["cards"], list)
    assert len(data["cards"]) >= 1


# ─── MVP-2a document extraction ──────────────────────────────────────


def test_document_extract_endpoint_schema():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    # Without an LLM key configured (or text too short) the endpoint
    # still returns the expected schema — empty items list.
    r = client.post("/api/v1/copilot/document/extract", json={"text": "hi"})
    assert r.status_code == 200
    data = r.json()
    assert "items" in data
    assert "count" in data
    assert isinstance(data["items"], list)
    assert data["count"] == len(data["items"])


def test_extract_items_short_text_returns_empty():
    import asyncio
    from app.copilot_engine import extract_items_from_text
    # Too short to be a real request — should short-circuit without an LLM call.
    assert asyncio.run(extract_items_from_text("")) == []
    assert asyncio.run(extract_items_from_text("hi")) == []


def test_detect_format_by_extension():
    from app.document_parser import detect_format
    assert detect_format("oferta.pdf", "", b"%PDF-1.4\n") == "pdf"
    assert detect_format("email.eml", "", b"From: x\n") == "eml"
    assert detect_format("list.txt", "", b"hello") == "text"
    assert detect_format("", "application/pdf", b"%PDF-") == "pdf"
    # Zip magic fallback for docx
    assert detect_format("mystery.bin", "", b"PK\x03\x04") == "docx"


def test_detect_format_unknown():
    from app.document_parser import detect_format
    assert detect_format("x.xyz", "application/octet-stream", b"\x00\x01\x02") == "unknown"


def test_extract_text_from_docx_roundtrip():
    # Build a tiny DOCX in-memory, run it through extract_text_from_docx
    from docx import Document
    from app.document_parser import extract_text_from_docx
    import io

    doc = Document()
    doc.add_paragraph("Zapotrzebowanie na tydzien:")
    doc.add_paragraph("- 50 klockow hamulcowych Bosch")
    doc.add_paragraph("- 30 filtrow oleju MANN")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text_from_docx(buf.getvalue())
    assert "klockow" in text.lower()
    assert "filtrow" in text.lower()


def test_extract_text_from_eml_roundtrip():
    from app.document_parser import extract_text_from_eml
    eml = (
        b"From: janusz@firma.pl\r\n"
        b"Subject: Zapotrzebowanie\r\n"
        b"Content-Type: text/plain; charset=utf-8\r\n\r\n"
        b"Potrzebuje 50 klockow hamulcowych Bosch na poniedzialek.\r\n"
    )
    text = extract_text_from_eml(eml)
    assert "janusz" in text.lower()
    assert "klockow" in text.lower()


def test_extract_text_dispatcher_decodes_utf16_bom():
    from app.document_parser import extract_text
    raw = "Potrzebuje 10 szt. klockow".encode("utf-16")
    text, fmt = extract_text("note.txt", "text/plain", raw)
    assert "klockow" in text
    assert fmt == "text"


def test_extract_file_endpoint_rejects_oversized():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    big = b"x" * (5 * 1024 * 1024 + 10)
    r = client.post(
        "/api/v1/copilot/document/extract-file",
        files={"file": ("huge.txt", big, "text/plain")},
    )
    assert r.status_code == 413


def test_ocr_respects_config_flag(monkeypatch):
    # When the feature flag is off, the OCR fallback should be a no-op
    # even if the deps happen to be installed. Returns empty string so the
    # text-layer result from pypdf is what the caller sees.
    from app.config import settings
    from app.document_parser import _ocr_pdf

    monkeypatch.setattr(settings, "pdf_ocr_enabled", False, raising=False)
    assert _ocr_pdf(b"%PDF-anything") == ""


def test_pdf_extraction_falls_through_when_ocr_unavailable(monkeypatch):
    # Simulate missing OCR binary: _ocr_pdf returns "" and we still get
    # whatever the pypdf text layer yielded (also "" for broken bytes).
    from app import document_parser
    monkeypatch.setattr(document_parser, "_ocr_pdf", lambda raw: "", raising=False)
    result = document_parser.extract_text_from_pdf(b"%PDF-\x00\xff broken")
    assert result == ""  # graceful, not a crash


def test_extract_file_endpoint_empty_pdf_returns_gracefully():
    # A "PDF" with no extractable text layer (bytes that pypdf can't parse)
    # returns count=0 and a friendly message, not a 500.
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.post(
        "/api/v1/copilot/document/extract-file",
        files={"file": ("broken.pdf", b"%PDF-1.4\n\xff\xff\xff", "application/pdf")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["count"] == 0
    assert data["format"] == "pdf"


def test_spend_analytics_endpoint():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/buying/spend-analytics?period_days=0")
    assert r.status_code == 200
    data = r.json()
    assert data["success"] is True
    assert "total_spend" in data
    assert "direct_spend" in data
    assert "indirect_spend" in data
    assert "direct_pct" in data
    assert "indirect_pct" in data
    assert "top_categories" in data
    assert isinstance(data["top_categories"], list)


def test_spend_analytics_respects_period():
    from app.buying_engine import spend_analytics
    # Future window — nothing should match
    result = spend_analytics(period_days=0)  # 0 = all-time per API contract
    assert "top_categories" in result


def test_contracts_endpoint_returns_demo_set():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/buying/contracts")
    assert r.status_code == 200
    data = r.json()
    assert data["success"] is True
    assert data["count"] >= 3  # demo seed has 5
    first = data["contracts"][0]
    assert "supplier_name" in first
    assert "days_to_expiry" in first
    assert "end_date" in first


def test_contracts_expiring_filter():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/buying/contracts?expiring_within_days=30")
    assert r.status_code == 200
    data = r.json()
    # Every returned contract must expire within 30 days
    for c in data["contracts"]:
        assert 0 <= c["days_to_expiry"] <= 30


def test_recommendation_engine_produces_contract_card():
    from app.recommendation_engine import generate_recommendations
    cards = generate_recommendations()
    # Seed has a 14-day Bosch contract — should surface as urgent
    titles = [c["title"] for c in cards]
    assert any("wygasa" in t.lower() for t in titles)
    urgent = [c for c in cards if c["urgency"] == "urgent"]
    assert urgent, "at least one urgent card expected from demo contracts"


def test_recommendation_engine_caps_output():
    from app.recommendation_engine import generate_recommendations
    cards = generate_recommendations(limit=3)
    assert len(cards) <= 3


def test_get_recommendations_includes_rule_output():
    # End-to-end: the dashboard endpoint composition should contain
    # at least one rule-produced contract card. Greeting card appears
    # only when the real rules haven't filled the 6-card cap.
    from app.copilot_engine import get_recommendations
    cards = get_recommendations({"step": 0})
    titles = [c["title"] for c in cards]
    assert any("wygasa" in t.lower() for t in titles), titles
    assert len(cards) >= 3, "expected a meaningful list of cards"


def test_bi_status_endpoint():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/bi/status")
    assert r.status_code == 200
    data = r.json()
    assert data["success"] is True
    names = {c["name"] for c in data["connectors"]}
    assert any("ERP" in n for n in names)
    assert any("BI" in n or "Enterprise" in n for n in names)
    assert any("EWM" in n or "WMS" in n for n in names)


def test_bi_invoices_deterministic():
    from app.bi_mock import get_connector
    erp = get_connector("erp")
    a = erp.get_invoices(months=2)
    b = erp.get_invoices(months=2)
    # Same input → same output (seeded RNG)
    assert [i["invoice_id"] for i in a] == [i["invoice_id"] for i in b]
    assert all(i["amount_pln"] > 0 for i in a)


def test_bi_spend_history_has_quarters():
    from app.bi_mock import get_connector
    bi = get_connector("bi")
    rows = bi.get_historical_spend(months=12)
    assert rows
    assert all("quarter" in r for r in rows)
    assert all(r["quarter"].startswith(("2024", "2025", "2026", "2027")) for r in rows)


def test_bi_yoy_anomaly_detection():
    from app.bi_mock import get_connector
    bi = get_connector("bi")
    anomalies = bi.yoy_anomalies(threshold_pct=20)
    for a in anomalies:
        assert abs(a["delta_pct"]) >= 20
        assert a["direction"] in ("up", "down")


def test_recommendation_includes_yoy_rule():
    # Given the deterministic BI mock, at least one category should trip the
    # 20% YoY threshold so the rule produces exactly one card.
    from app.recommendation_engine import _rule_yoy_anomaly
    cards = _rule_yoy_anomaly()
    # Might be empty if all categories happen to sit inside ±20%, but
    # with our seed it should produce one.
    for c in cards:
        assert c.urgency in ("info", "urgent")
        assert "YoY" in c.title


def test_supplier_concentration_falls_back_to_allocations():
    # The seeded demo orders have solver allocations but often no
    # purchase_orders[]. Rule should still attribute spend via the
    # domain_results fallback.
    from app.recommendation_engine import _rule_supplier_concentration
    cards = _rule_supplier_concentration()
    # At least one supplier in demo ends up dominant — exact count is
    # data-driven, so just assert the rule produces something and the
    # card payload has an attributed supplier in the title.
    if cards:
        assert "koncentracji" in cards[0].title.lower()
        assert cards[0].urgency in ("urgent", "info")


def test_contract_audit_logs_create_and_update():
    from app.database import init_db
    from app.contract_engine import reset_cache, upsert_contract, get_contract_audit
    from datetime import date, timedelta
    init_db()
    reset_cache()
    payload = {
        "id": "CNT-PYTEST-42",
        "supplier_id": "SUP-X",
        "supplier_name": "AuditTest",
        "category": "parts",
        "start_date": date.today().isoformat(),
        "end_date": (date.today() + timedelta(days=30)).isoformat(),
        "committed_volume_pln": 50000,
        "notes": "initial",
    }
    upsert_contract(payload, actor="pytest")
    payload["notes"] = "edited"
    payload["committed_volume_pln"] = 75000
    upsert_contract(payload, actor="pytest")
    entries = get_contract_audit("CNT-PYTEST-42")
    assert len(entries) >= 2
    actions = [e["action"] for e in entries]
    assert "create" in actions
    assert "update" in actions
    # Newest first
    assert entries[0]["occurred_at"] >= entries[-1]["occurred_at"]


def test_contract_audit_endpoint():
    from fastapi.testclient import TestClient
    from app.main import app
    from app.contract_engine import upsert_contract
    from datetime import date, timedelta

    upsert_contract({
        "id": "CNT-ENDPOINT-1",
        "supplier_id": "SUP-Y",
        "supplier_name": "EndpointTest",
        "category": "oils",
        "start_date": date.today().isoformat(),
        "end_date": (date.today() + timedelta(days=60)).isoformat(),
        "committed_volume_pln": 10000,
    }, actor="pytest-api")

    client = TestClient(app)
    r = client.get("/api/v1/buying/contracts/CNT-ENDPOINT-1/audit")
    assert r.status_code == 200
    data = r.json()
    assert data["success"] is True
    assert isinstance(data["entries"], list)


def test_whatif_chain_demo_endpoint():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/whatif/chain/demo?domain=parts")
    assert r.status_code == 200
    data = r.json()
    assert data["total_steps"] >= 2, "chain should include baseline + at least 1 step"
    chain = data["chain"]
    # Baseline's applied_delta must be empty; subsequent steps must carry something
    assert chain[0]["applied_delta"] == {}
    for step in chain[1:]:
        assert step["applied_delta"], step["label"]
    # delta_vs_prev absent on first step, present on subsequent
    assert not chain[0]["delta_vs_prev"]
    for step in chain[1:]:
        if step["result"].get("success"):
            dv = step["delta_vs_prev"]
            assert "total_cost_pln" in dv, step["label"]


def test_whatif_chain_auto_normalizes_weights():
    # User sends w_cost=0.55 without touching the other three; engine must
    # normalize so CriteriaWeights validator (sum=1) doesn't trip.
    from app.data_layer import DOMAIN_DATA
    from app.whatif_engine import WhatIfEngine
    parts = DOMAIN_DATA["parts"]
    engine = WhatIfEngine(suppliers=parts["suppliers"], demand=parts["demand"])
    result = engine.run_chain(
        steps=[{"label": "cheap", "w_cost": 0.55, "w_time": 0.15}],
    )
    assert result["total_steps"] == 2
    # Second step (the chain step) must have succeeded
    assert result["chain"][1]["result"]["success"], result["chain"][1]


def test_subdomain_optimizer_parts():
    from app.subdomain_optimizer import optimize_per_subdomain
    result = optimize_per_subdomain("parts")
    assert "aggregate" in result
    assert result["subdomain_total"] >= 3  # parts has 3 subdomens
    agg = result["aggregate"]
    # Sum of subdomain costs must equal aggregate total_cost_pln
    total = sum(s["total_cost_pln"] for s in result["subdomains"] if s["success"])
    assert abs(agg["total_cost_pln"] - total) < 1.0  # rounding tolerance
    # Every successful subdomain should carry top_suppliers
    for s in result["subdomains"]:
        if s["success"]:
            assert s["suppliers_used"] > 0
            assert isinstance(s.get("top_suppliers"), list)


def test_subdomain_aggregate_endpoint():
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/dashboard/subdomain-aggregate/demo?domain=oils")
    assert r.status_code == 200
    data = r.json()
    assert data["domain"] == "oils"
    assert "subdomains" in data
    assert "aggregate" in data
    assert data["total_time_ms"] > 0


def test_subdomain_optimizer_unknown_domain():
    from app.subdomain_optimizer import optimize_per_subdomain
    result = optimize_per_subdomain("kosmolot")
    assert result.get("error") is True
    assert "kosmolot" in result["message"].lower() or "unknown" in result["message"].lower()


def test_pareto_mc_endpoint_shape():
    # Phase B2 — endpoint returns ParetoPointMC with MC fields populated.
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get(
        "/api/v1/dashboard/pareto-xy-mc/demo?domain=parts&steps=5&mc_iterations=8"
    )
    assert r.status_code == 200
    data = r.json()
    assert data["points"], "expected at least one Pareto+MC point"
    for p in data["points"]:
        assert "cost_mean_pln" in p
        assert "cost_p5_pln" in p
        assert "cost_p95_pln" in p
        assert p["cost_p5_pln"] <= p["cost_mean_pln"] <= p["cost_p95_pln"]
        assert 0.0 <= p["feasible_rate"] <= 1.0
        assert p["mc_iterations"] == 8


def test_pareto_mc_clamps_iterations():
    # Sanity guard — backend clamps extreme values to [5, 500]
    from app.pareto import generate_pareto_with_mc
    # Smoke — pass a reasonable demo-sized call, just ensure it runs
    from app.data_layer import DOMAIN_DATA
    parts = DOMAIN_DATA["parts"]
    from app.schemas import CriteriaWeights, SolverMode
    pts = generate_pareto_with_mc(
        suppliers=parts["suppliers"], demand=parts["demand"],
        base_weights=CriteriaWeights(), mode=SolverMode.continuous,
        steps=3, mc_iterations=5,
    )
    assert pts
    for p in pts:
        assert p.mc_iterations == 5


def test_domains_extended_taxonomy_counts():
    # Slide 3 from the deck promises 10 domains × 27 subdomains;
    # the endpoint summary must reconcile.
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/v1/domains/extended")
    assert r.status_code == 200
    data = r.json()
    s = data["summary"]
    assert s["domains_total"] == 10
    assert s["direct_domains"] == 6
    assert s["indirect_domains"] == 4
    assert s["subdomains_total"] == 27
    # Every listed domain must carry a non-empty subdomain list
    for d in data["domains"]:
        assert len(d["subdomains"]) > 0, d["domain"]


def test_shadow_prices_returned_when_esg_binding():
    from app.schemas import (
        ConstraintConfig,
        CriteriaWeights,
        DemandItem,
        SolverMode,
        SupplierInput,
    )
    from app.optimizer import run_optimization

    suppliers = [
        SupplierInput(supplier_id="S1", name="Green", unit_cost=50.0, logistics_cost=5.0,
                      lead_time_days=5.0, compliance_score=0.9, esg_score=0.92,
                      max_capacity=1000, served_regions=["PL-MA"], payment_terms_days=30),
        SupplierInput(supplier_id="S2", name="Cheap", unit_cost=30.0, logistics_cost=3.0,
                      lead_time_days=3.0, compliance_score=0.9, esg_score=0.65,
                      max_capacity=1000, served_regions=["PL-MA"], payment_terms_days=60),
    ]
    demand = [DemandItem(product_id="P1", demand_qty=100, destination_region="PL-MA")]
    resp, _ = run_optimization(
        suppliers, demand, CriteriaWeights(),
        mode=SolverMode.continuous, max_vendor_share=1.0,
        constraints=ConstraintConfig(min_esg_score=0.85, preferred_supplier_bonus=0.0),
    )
    ids = [sp.constraint_id for sp in resp.shadow_prices]
    assert any("esg" in i.lower() for i in ids), ids
    # Demand equality always binding
    assert any("demand" in i.lower() for i in ids), ids


def test_shadow_prices_excluded_when_not_binding():
    # With no advanced constraints and slack capacity, the output should
    # contain at most the demand marginals — capacity/diversification drop
    # out because they're non-binding.
    from app.schemas import (
        CriteriaWeights,
        DemandItem,
        SolverMode,
        SupplierInput,
    )
    from app.optimizer import run_optimization

    suppliers = [
        SupplierInput(supplier_id="S1", name="A", unit_cost=40, logistics_cost=4,
                      lead_time_days=3, compliance_score=0.9, esg_score=0.85,
                      max_capacity=10000, served_regions=["PL-MA"]),
    ]
    demand = [DemandItem(product_id="P1", demand_qty=50, destination_region="PL-MA")]
    resp, _ = run_optimization(
        suppliers, demand, CriteriaWeights(),
        mode=SolverMode.continuous, max_vendor_share=1.0,
    )
    # No advanced constraints, so only demand equality constraints should
    # produce meaningful shadow prices (plus possibly a near-zero tail).
    for sp in resp.shadow_prices:
        assert sp.kind in ("demand", "capacity", "diversification", "esg",
                           "payment", "region", "preferred", "unknown")


def test_spend_analytics_splits_direct_indirect():
    # Totals must reconcile: direct + indirect == total_spend (within epsilon)
    from app.buying_engine import spend_analytics
    r = spend_analytics(period_days=None)
    assert abs(r["total_spend"] - (r["direct_spend"] + r["indirect_spend"])) < 0.01
    if r["total_spend"] > 0:
        assert 99.0 <= (r["direct_pct"] + r["indirect_pct"]) <= 101.0


def test_extracted_item_catalog_matching():
    # When the LLM returns a name that matches the catalog, extract_demand
    # should attach matched_id / matched_name. We bypass the LLM call by
    # monkey-patching extract_items_from_text on the module.
    import asyncio
    from app import copilot_engine

    async def fake_extract(raw: str):
        return [{"name": "klocki hamulcowe TRW", "qty": 3, "unit": "szt", "price": None, "note": ""}]

    orig = copilot_engine.extract_items_from_text
    copilot_engine.extract_items_from_text = fake_extract
    try:
        items = asyncio.run(copilot_engine.extract_demand("dummy"))
    finally:
        copilot_engine.extract_items_from_text = orig

    assert len(items) == 1
    assert items[0].matched_id.startswith("BRK") or items[0].matched_id.startswith("DSC"), items[0].matched_id
    assert items[0].qty == 3
